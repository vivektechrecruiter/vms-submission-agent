from docx import Document


def read_docx(file_path):

    doc = Document(file_path)

    text = []

    # ==========================
    # Headers
    # ==========================

    for section in doc.sections:

        header = section.header

        for paragraph in header.paragraphs:

            if paragraph.text.strip():

                text.append(
                    paragraph.text
                )

    # ==========================
    # Body
    # ==========================

    for paragraph in doc.paragraphs:

        if paragraph.text.strip():

            text.append(
                paragraph.text
            )

    return "\n".join(text)