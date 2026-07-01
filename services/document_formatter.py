from docx import Document
from docx.shared import Inches, Pt
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


SECTION_HEADERS = {
    "PROFESSIONAL SUMMARY",
    "SKILLS",
    "PROFESSIONAL EXPERIENCE",
    "EDUCATION",
    "CERTIFICATIONS"
}


def format_resume_document(doc_path):

    doc = Document(doc_path)

    section = doc.sections[0]

    # Narrow Margins
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

    # Portrait Orientation
    section.orientation = WD_ORIENT.PORTRAIT

    for paragraph in doc.paragraphs:

        text = paragraph.text.strip()

        # Default formatting
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

        for run in paragraph.runs:
            run.font.name = "Arial"
            run.font.size = Pt(10)

        # Candidate Name (first paragraph)
        if text and text == doc.paragraphs[0].text.strip():

            for run in paragraph.runs:
                run.bold = True
                run.font.name = "Arial"
                run.font.size = Pt(10)

            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        # Section Headers
        elif text.upper() in SECTION_HEADERS:

            for run in paragraph.runs:
                run.bold = True
                run.font.name = "Arial"
                run.font.size = Pt(10)

            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        # Responsibilities / Subsection Headers
        elif text.endswith(":"):

            for run in paragraph.runs:
                run.bold = True
                run.font.name = "Arial"
                run.font.size = Pt(10)

            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    doc.save(doc_path)