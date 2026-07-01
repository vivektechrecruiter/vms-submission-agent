from docx import Document
from docx.shared import Pt, Inches


def add_row(table, label, value):

    row = table.add_row().cells

    row[0].text = label
    row[1].text = value

    for cell in row:

        for paragraph in cell.paragraphs:

            for run in paragraph.runs:
                run.font.name = "Arial"
                run.font.size = Pt(10)


def generate_submission_sheet(
    candidate_data,
    recruiter_data,
    output_path
):

    doc = Document()

    section = doc.sections[0]

    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

    title = doc.add_paragraph()

    run = title.add_run(
        "Candidate Submittal Cover Page"
    )

    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(12)

    table = doc.add_table(
        rows=0,
        cols=2
    )

    table.style = "Table Grid"

    add_row(
        table,
        "Candidate Name",
        candidate_data.get(
            "candidate_name",
            ""
        )
    )

    add_row(
        table,
        "Phone Number",
        candidate_data.get(
            "phone",
            ""
        )
    )

    add_row(
        table,
        "Email Address",
        candidate_data.get(
            "email",
            ""
        )
    )

    add_row(
        table,
        "LinkedIn",
        candidate_data.get(
            "linkedin",
            ""
        )
    )

    add_row(
        table,
        "Current Location",
        candidate_data.get(
            "location",
            ""
        )
    )

    add_row(
        table,
        "Willing To Relocate",
        recruiter_data.get(
            "willing_to_relocate",
            ""
        )
    )

    add_row(
        table,
        "Former TCS Employee",
        recruiter_data.get(
            "former_tcs_employee",
            ""
        )
    )

    add_row(
        table,
        "Interview Availability",
        recruiter_data.get(
            "interview_availability",
            ""
        )
    )

    add_row(
        table,
        "General Availability",
        recruiter_data.get(
            "general_availability",
            ""
        )
    )

    add_row(
        table,
        "Availability To Start",
        recruiter_data.get(
            "availability_to_start",
            ""
        )
    )

    doc.save(output_path)