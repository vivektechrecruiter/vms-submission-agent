from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


def _format_paragraph(p, alignment=WD_PARAGRAPH_ALIGNMENT.LEFT):
    p.alignment = alignment
    p_format = p.paragraph_format
    p_format.space_before = Pt(0)
    p_format.space_after = Pt(0)
    p_format.line_spacing = 1
    return p


def _add_paragraph(doc, text="", bold=False, alignment=WD_PARAGRAPH_ALIGNMENT.LEFT):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(10)
    _format_paragraph(p, alignment=alignment)
    return p


def generate_resume_docx(candidate_data, output_path):

    doc = Document()

    # ==========================
    # Page Setup
    # ==========================

    section = doc.sections[0]
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

    

    # ==========================
    # PROFESSIONAL SUMMARY
    # ==========================

    _add_paragraph(
        doc,
        "PROFESSIONAL SUMMARY",
        bold=True,
        alignment=WD_PARAGRAPH_ALIGNMENT.LEFT
    )

    summary_list = candidate_data.get("professional_summary", [])
    for summary in summary_list:
        _add_paragraph(
            doc,
            f"• {summary}",
            alignment=WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        )

    # Blank line between major sections
    _add_paragraph(doc)

    # ==========================
    # SKILLS
    # ==========================

    _add_paragraph(
        doc,
        "SKILLS",
        bold=True,
        alignment=WD_PARAGRAPH_ALIGNMENT.LEFT
    )

    skills_list = candidate_data.get("skills_summary", [])
    for skill in skills_list:
        _add_paragraph(
            doc,
            skill,
            alignment=WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        )

    _add_paragraph(doc)

    # ==========================
    # PROFESSIONAL EXPERIENCE
    # ==========================

    _add_paragraph(
        doc,
        "PROFESSIONAL EXPERIENCE",
        bold=True,
        alignment=WD_PARAGRAPH_ALIGNMENT.LEFT
    )

    employment_history = candidate_data.get("employment_history", [])
    for idx, job in enumerate(employment_history):
        company = job.get("company", "")
        location = job.get("location", "")
        start_date = job.get("start_date", "")
        end_date = job.get("end_date", "")
        title = job.get("title", "")

        company_line = f"{company}, {location}"
        if start_date or end_date:
            date_line = f"{start_date} - {end_date}".strip()
            company_line = f"{company}, {location} {date_line}"
        _add_paragraph(
            doc,
            company_line,
            bold=True,
            alignment=WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        )

        if title:
            _add_paragraph(
                doc,
                title,
                bold=True,
                alignment=WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            )

        _add_paragraph(
            doc,
            "Responsibilities:",
            bold=True,
            alignment=WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        )

        for responsibility in job.get("responsibilities", []):
            _add_paragraph(
                doc,
                f"• {responsibility}",
                alignment=WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            )

        for subsection in job.get("subsections", []):
            subsection_name = subsection.get("name", "")
            if subsection_name:
                _add_paragraph(
                    doc,
                    f"{subsection_name}:",
                    bold=True,
                    alignment=WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                )

            for responsibility in subsection.get("responsibilities", []):
                _add_paragraph(
                    doc,
                    f"• {responsibility}",
                    alignment=WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                )

        # Add exactly one blank line between companies (but not after the last company)
        if idx < len(employment_history) - 1:
            _add_paragraph(doc)

    if candidate_data.get("education"):
        _add_paragraph(doc)
        _add_paragraph(
            doc,
            "EDUCATION",
            bold=True,
            alignment=WD_PARAGRAPH_ALIGNMENT.LEFT
        )
        for education in candidate_data.get("education", []):
            _add_paragraph(
                doc,
                education,
                alignment=WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            )

    if candidate_data.get("certifications"):
        _add_paragraph(doc)
        _add_paragraph(
            doc,
            "CERTIFICATIONS",
            bold=True,
            alignment=WD_PARAGRAPH_ALIGNMENT.LEFT
        )
        for certification in candidate_data.get("certifications", []):
            _add_paragraph(
                doc,
                certification,
                alignment=WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            )

    doc.save(output_path)
