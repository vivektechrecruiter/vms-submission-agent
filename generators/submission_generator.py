from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import json


def format_cell(cell):
    """Format cell text to Arial 10."""
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.name = "Arial"
            run.font.size = Pt(10)


def add_row(table, label, value):
    """Add a 2-column row with label and value."""
    row = table.add_row().cells
    
    # Left cell: label (bold) with colon if not already punctuated
    if label.endswith(":") or label.endswith("?"):
        row[0].text = label
    else:
        row[0].text = label + ":"
    
    for paragraph in row[0].paragraphs:
        for run in paragraph.runs:
            run.font.name = "Arial"
            run.font.size = Pt(10)
            run.bold = True
    
    # Right cell: value
    row[1].text = value
    for paragraph in row[1].paragraphs:
        for run in paragraph.runs:
            run.font.name = "Arial"
            run.font.size = Pt(10)


def add_skills_row(table, skill="", years="", experience=""):
    """Add a 3-column skills row."""
    row = table.add_row().cells
    row[0].text = skill
    row[1].text = years
    row[2].text = experience

    for cell in row:
        format_cell(cell)


def generate_submission_sheet(
    candidate_data,
    recruiter_data,
    skill_match_data,
    output_path
):

    doc = Document(
    "templates/tcs_submission_template.docx"
)
    table = doc.tables[0]

    section = doc.sections[0]
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

    # ==========================
    # Title
    # ==========================
    title = doc.add_paragraph()
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = title.add_run("Candidate Submittal Cover Page")
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(14)

    # ==========================
    # Candidate Info Table (2 columns)
    # ==========================
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    
    # Set fixed column widths
    table.columns[0].width = Inches(2.5)
    table.columns[1].width = Inches(3.5)

    add_row(
        table,
        "Candidate Name",
        candidate_data.get("candidate_name", "")
    )

    add_row(
        table,
        "Phone Number",
        candidate_data.get("phone", "")
    )

    add_row(
        table,
        "Email Address",
        candidate_data.get("email", "")
    )

    add_row(
        table,
        "LinkedIn",
        candidate_data.get("linkedin", "")
    )

    add_row(
        table,
        "Current Location",
        candidate_data.get("location", "")
    )

    add_row(
        table,
        "Willing to Relocate?",
        recruiter_data.get("willing_to_relocate", "")
    )

    add_row(
        table,
        "Former TCS employee / contractor?\n(Please specify employment type and when)",
        recruiter_data.get("former_tcs_employee", "")
    )

    add_row(
        table,
        "General Interview Availability\n(INCLUDE TIMEZONE)",
        recruiter_data.get("interview_availability", "")
    )

    add_row(
        table,
        "Candidate's General Availability",
        recruiter_data.get("general_availability", "")
    )

    add_row(
        table,
        "Availability to Start",
        recruiter_data.get("availability_to_start", "")
    )

    # Blank line
    doc.add_paragraph()

    # ==========================
    # Relevant Skills Heading
    # ==========================
    skills_heading = doc.add_paragraph()
    run = skills_heading.add_run("Relevant Skills")
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(10)

    # ==========================
    # Skills Table (3 columns with 5 empty rows)
    # ==========================
    skills_table = doc.add_table(rows=1, cols=3)
    skills_table.style = "Table Grid"
    
    # Set fixed column widths
    skills_table.columns[0].width = Inches(2.0)
    skills_table.columns[1].width = Inches(1.5)
    skills_table.columns[2].width = Inches(2.5)

    # Header row
    header_cells = skills_table.rows[0].cells
    header_cells[0].text = "Mandatory Skills (As listed in JD)"
    header_cells[1].text = "# Of Years' Experience"
    header_cells[2].text = "Candidate's Relevant Hands-On Experience"

    for cell in header_cells:
        format_cell(cell)

        # Populate submission skills

    for skill in skill_match_data.get(
        "submission_skills",
        []
    ):

        row_cells = (
            skills_table
            .add_row()
            .cells
        )

        row_cells[0].text = skill.get(
            "skill",
            ""
        )

        row_cells[1].text = skill.get(
            "years_experience",
            ""
        )

        row_cells[2].text = skill.get(
            "relevant_experience",
            ""
        )

        for cell in row_cells:
            format_cell(cell)

    # Blank line
    doc.add_paragraph()

    # ==========================
    # Resume Placeholder Heading
    # ==========================
   

    doc.save(output_path)