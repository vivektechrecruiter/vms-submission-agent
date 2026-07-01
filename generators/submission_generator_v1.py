from docx import Document
from docx.shared import Pt, Inches


def generate_submission_sheet(
    candidate_data,
    recruiter_data,
    output_path
):

    doc = Document()

    section = doc.sections[0]

    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

    title = doc.add_paragraph()

    run = title.add_run(
        "Candidate Submittal Cover Page"
    )

    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(12)

    doc.add_paragraph(
        f"Candidate Name: {candidate_data.get('candidate_name', '')}"
    )

    doc.add_paragraph(
        f"Phone Number: {candidate_data.get('phone', '')}"
    )

    doc.add_paragraph(
        f"Email Address: {candidate_data.get('email', '')}"
    )

    doc.add_paragraph(
        f"LinkedIn: {candidate_data.get('linkedin', '')}"
    )

    doc.add_paragraph(
        f"Current Location: {candidate_data.get('location', '')}"
    )

    doc.add_paragraph(
        f"Willing To Relocate: "
        f"{recruiter_data.get('willing_to_relocate', '')}"
    )

    doc.add_paragraph(
        f"Former TCS Employee: "
        f"{recruiter_data.get('former_tcs_employee', '')}"
    )

    doc.add_paragraph(
        f"Interview Availability: "
        f"{recruiter_data.get('interview_availability', '')}"
    )

    doc.add_paragraph(
        f"General Availability: "
        f"{recruiter_data.get('general_availability', '')}"
    )

    doc.add_paragraph(
        f"Availability To Start: "
        f"{recruiter_data.get('availability_to_start', '')}"
    )

    doc.save(output_path)